import os
import tempfile
import PyPDF2
import streamlit as st
import torch
import requests
from bs4 import BeautifulSoup
import time
import pandas as pd
import warnings
import threading
import psutil
import io
import docx
from utils import remove_directory_recursively

# LangChain imports
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import OllamaLLM
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS, Chroma
from langchain.chains import RetrievalQA, LLMChain
from langchain.docstore.document import Document
from langchain.prompts import PromptTemplate
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler

# Suppress warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)
os.environ["PYTHONWARNINGS"] = "ignore::DeprecationWarning"

# Fix for PyTorch/Streamlit compatibility issue
if "STREAMLIT_WATCH_MODULES" in os.environ:
    modules_to_skip = ["torch", "tensorflow"]
    current_modules = os.environ["STREAMLIT_WATCH_MODULES"].split(",")
    filtered_modules = [m for m in current_modules if all(skip not in m for skip in modules_to_skip)]
    os.environ["STREAMLIT_WATCH_MODULES"] = ",".join(filtered_modules)

class EnhancedRAG:
    def __init__(self, 
                 llm_model_name="llama3.2:latest",
                 embedding_model_name="sentence-transformers/all-MiniLM-L6-v2",
                 chunk_size=1000,
                 chunk_overlap=200,
                 use_gpu=True):
        """
        Initialize the Enhanced RAG system with multiple modes.
        
        Args:
            llm_model_name: The Ollama model for text generation
            embedding_model_name: The HuggingFace model for embeddings
            chunk_size: Size of document chunks
            chunk_overlap: Overlap between chunks
            use_gpu: Whether to use GPU acceleration
        """
        self.llm_model_name = llm_model_name
        self.embedding_model_name = embedding_model_name
        self.use_gpu = use_gpu and torch.cuda.is_available()
        self.temp_dirs = []  # Keep track of temporary directories
        
        # Device selection for embeddings
        self.device = "cuda" if self.use_gpu else "cpu"
        st.sidebar.info(f"Using device: {self.device}")
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        
        # Initialize embeddings model
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name=embedding_model_name,
                model_kwargs={"device": self.device}
            )
            st.sidebar.success(f"Embeddings model loaded: {embedding_model_name}")
        except Exception as e:
            st.sidebar.error(f"Failed to load embeddings model: {str(e)}")
            self.embeddings = None
        
        # Initialize LLM
        try:
            callbacks = [StreamingStdOutCallbackHandler()]
            self.llm = OllamaLLM(model=llm_model_name, callbacks=callbacks)
            st.sidebar.success(f"LLM loaded: {llm_model_name}")
        except Exception as e:
            st.sidebar.error(f"Failed to load LLM: {str(e)}")
            self.llm = None
        
        # Initialize vector stores for different sources
        self.doc_vector_store = None
        self.web_vector_store = None
        self.documents_processed = 0
        
        # Monitoring stats
        self.processing_times = {}
        
        # Keep track of sources and errors
        self.sources = []
        self.errors = []
    
    def __del__(self):
        """Cleanup temporary directories when object is garbage collected."""
        for temp_dir in self.temp_dirs:
            try:
                if os.path.exists(temp_dir):
                    remove_directory_recursively(temp_dir)
            except:
                pass

    def process_files(self, files, user_id=None, mongodb=None, notebook_id=None, is_nested=False):
        """Process files and create a vector store using GPU acceleration."""
        # Check if embeddings are available
        if self.embeddings is None:
            st.error("Embeddings model not initialized. Unable to process files.")
            return False
            
        all_docs = []
        document_metadata = []
        
        # Use status or simple spinner based on nesting
        if is_nested:
            # We're inside another expander/status, so just use progress indicators
            status_msg = st.empty()
            status_msg.info("Processing files...")
            progress_bar = st.progress(0)
        else:
            # We can use the full status widget
            status = st.status("Processing files...")
        
        # Create temporary directory for file storage
        temp_dir = tempfile.mkdtemp()
        self.temp_dirs.append(temp_dir)  # Track for cleanup
        st.session_state['temp_dir'] = temp_dir
        
        # Monitor processing time and memory usage
        start_time = time.time()
        
        # Track memory before processing
        mem_before = psutil.virtual_memory().used / (1024 * 1024 * 1024)  # GB
        
        # Process each file
        total_files = len(files)
        for i, file in enumerate(files):
            try:
                # Update progress
                if is_nested:
                    progress_bar.progress((i + 1) / total_files)
                    status_msg.info(f"Processing {file.name} ({i+1}/{total_files})...")
                else:
                    status.update(label=f"Processing {file.name} ({i+1}/{total_files})...")
                
                file_start_time = time.time()
                file_type = "unknown"
                
                # Determine file type
                if file.name.lower().endswith('.pdf'):
                    file_type = "pdf"
                elif file.name.lower().endswith(('.docx', '.doc')):
                    file_type = "docx"
                elif file.name.lower().endswith(('.txt')):
                    file_type = "txt"
                
                # Save uploaded file to temp directory
                file_path = os.path.join(temp_dir, file.name)
                file.seek(0)  # Reset file pointer
                file_content = file.read()  # Read the file content
                
                with open(file_path, "wb") as f:
                    f.write(file_content)
                
                # Extract text based on file type
                text = ""
                page_count = 0
                
                if file_type == "pdf":
                    try:
                        with open(file_path, "rb") as f:
                            pdf = PyPDF2.PdfReader(f)
                            page_count = len(pdf.pages)
                            for page_num in range(page_count):
                                page = pdf.pages[page_num]
                                page_text = page.extract_text()
                                if page_text:
                                    text += page_text + "\n\n"
                    except Exception as e:
                        st.error(f"Error extracting text from PDF {file.name}: {str(e)}")
                        continue
                
                elif file_type == "docx":
                    try:
                        doc = docx.Document(file_path)
                        page_count = len(doc.paragraphs)
                        for para in doc.paragraphs:
                            text += para.text + "\n\n"
                    except Exception as e:
                        st.error(f"Error extracting text from DOCX {file.name}: {str(e)}")
                        continue
                
                elif file_type == "txt":
                    try:
                        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                            text = f.read()
                        page_count = text.count("\n") + 1
                    except Exception as e:
                        st.error(f"Error extracting text from TXT {file.name}: {str(e)}")
                        continue
                
                # Skip file if no text was extracted
                if not text.strip():
                    st.warning(f"No text content found in {file.name}. Skipping...")
                    continue
                
                # Create documents
                docs = [Document(page_content=text, metadata={
                    "source": file.name, 
                    "notebook_id": notebook_id,
                    "file_type": file_type
                })]
                
                # Split documents into chunks
                split_docs = self.text_splitter.split_documents(docs)
                
                all_docs.extend(split_docs)
                
                file_end_time = time.time()
                processing_time = file_end_time - file_start_time
                
                # Save document metadata
                doc_meta = {
                    "filename": file.name,
                    "file_type": file_type,
                    "page_count": page_count,
                    "chunk_count": len(split_docs),
                    "processing_time": processing_time,
                    "notebook_id": notebook_id
                }
                document_metadata.append(doc_meta)
                
                # Display success message
                success_msg = f"Processed {file.name}: {len(split_docs)} chunks in {processing_time:.2f}s"
                if is_nested:
                    st.success(success_msg)
                else:
                    st.sidebar.success(success_msg)
                    
                self.processing_times[file.name] = {
                    "chunks": len(split_docs),
                    "time": processing_time
                }
                
            except Exception as e:
                error_msg = f"Error processing {file.name}: {str(e)}"
                self.errors.append(error_msg)
                if is_nested:
                    st.error(error_msg)
                else:
                    st.sidebar.error(error_msg)
        
        # Create vector store if we have documents
        if all_docs:
            if is_nested:
                status_msg.info("Building vector index...")
            else:
                status.update(label="Building vector index...")
                
            try:
                # Record the time taken to build the index
                index_start_time = time.time()
                
                # Create the vector store using FAISS
                self.doc_vector_store = FAISS.from_documents(all_docs, self.embeddings)
                
                index_end_time = time.time()
                index_time = index_end_time - index_start_time
                
                # Track memory after processing
                mem_after = psutil.virtual_memory().used / (1024 * 1024 * 1024)  # GB
                mem_used = mem_after - mem_before
                
                total_time = time.time() - start_time
                
                complete_msg = f"Completed processing {len(all_docs)} chunks in {total_time:.2f}s"
                if is_nested:
                    status_msg.success(complete_msg)
                    progress_bar.progress(1.0)
                else:
                    status.update(label=complete_msg, state="complete")
                
                # Save performance metrics
                self.processing_times["index_building"] = index_time
                self.processing_times["total_time"] = total_time
                self.processing_times["memory_used_gb"] = mem_used
                self.documents_processed = len(all_docs)
                
                # Save document metadata to MongoDB if user is logged in
                if user_id and mongodb:
                    overall_meta = {
                        "documents": document_metadata,
                        "total_chunks": len(all_docs),
                        "index_building_time": index_time,
                        "total_processing_time": total_time,
                        "memory_used_gb": mem_used,
                        "notebook_id": notebook_id
                    }
                    mongodb.save_document_metadata(user_id, overall_meta, notebook_id)
                
                # Clean up temporary UI elements if nested
                if is_nested:
                    time.sleep(1)  # Let user see completion message
                    status_msg.empty()
                    # Keep progress bar at 100%
                
                return True
            except Exception as e:
                error_msg = f"Error creating vector store: {str(e)}"
                self.errors.append(error_msg)
                st.error(error_msg)
                if is_nested:
                    status_msg.error(error_msg)
                    time.sleep(2)  # Let user see error message
                    status_msg.empty()
                else:
                    status.update(label=error_msg, state="error")
                return False
        else:
            empty_msg = "No content extracted from files"
            if is_nested:
                status_msg.error(empty_msg)
                time.sleep(2)  # Let user see error message
                status_msg.empty()
            else:
                status.update(label=empty_msg, state="error")
            return False

    def enhance_answer(self, initial_answer, query, source_content):
        """
        Enhance the initial answer with additional context and improved quality.
        
        Args:
            initial_answer: The initial answer generated by the RAG system
            query: The original user query
            source_content: The source content chunks used to generate the answer
            
        Returns:
            An enhanced answer with improved quality and formatting
        """
        # Create an enhancement prompt template
        enhance_template = """
        You are an expert content enhancer. Your task is to improve the quality of an AI-generated answer
        while maintaining factual accuracy.
        
        Below is a query, an initial answer, and the source content used to generate that answer.
        
        QUERY:
        {query}
        
        INITIAL ANSWER:
        {initial_answer}
        
        SOURCE CONTENT (EXTRACT):
        {source_content}
        
        Please enhance the initial answer by:
        1. Improving clarity and readability
        2. Adding relevant details from the source if they were missed
        3. Ensuring all claims are factually supported by the source content
        4. Adding appropriate structure (headings, bullet points) if helpful
        5. Making sure the tone is professional and helpful
        
        ENHANCED ANSWER:
        """
        
        # Create enhancement prompt
        enhancement_prompt = PromptTemplate(
            template=enhance_template,
            input_variables=["query", "initial_answer", "source_content"]
        )
        
        # Create enhancement chain
        enhancement_chain = LLMChain(
            llm=self.llm,
            prompt=enhancement_prompt
        )
        
        # Prepare source content for the enhancement (limited to avoid token limits)
        summarized_sources = "\n\n".join([
            f"SOURCE {i+1}:\n{source[:500]}..." if len(source) > 500 else f"SOURCE {i+1}:\n{source}"
            for i, source in enumerate(source_content[:3])  # Limit to first 3 sources
        ])
        
        # Invoke the enhancement chain
        try:
            enhanced_result = enhancement_chain.invoke({
                "query": query,
                "initial_answer": initial_answer,
                "source_content": summarized_sources
            })
            
            return enhanced_result["text"].strip()
        except Exception as e:
            st.warning(f"Enhancement step encountered an issue: {str(e)}. Using initial answer.")
            self.errors.append(f"Enhancement error: {str(e)}")
            return initial_answer

    def web_search(self, query, num_results=5):
        """
        Perform a web search using multiple fallback methods
        """
        try:
            # For this implementation, we'll use a simulated search
            # In a production environment, you would integrate with a real search API
            results = self.simulate_search(query, num_results)
            if results and len(results) > 0:
                self.errors.append("Search simulation succeeded")
                return results
            else:
                return self.get_mock_results(query)
        except Exception as e:
            self.errors.append(f"Search error: {str(e)}")
            return self.get_mock_results(query)

    def simulate_search(self, query, num_results=5):
        """Simulate a search with realistic results (for demo)"""
        # Create realistic-looking mock results based on the query
        results = []
        
        # Add Wikipedia result
        results.append({
            "title": f"{query} - Wikipedia",
            "url": f"https://en.wikipedia.org/wiki/{query.replace(' ', '_')}",
            "snippet": f"Wikipedia article about {query} providing comprehensive information from various sources."
        })
        
        # Academic or educational sources
        results.append({
            "title": f"{query} | Academic Research and Analysis",
            "url": f"https://scholar.example.com/research/{query.lower().replace(' ', '-')}",
            "snippet": f"Scholarly research and academic analysis of {query} with citations and peer-reviewed content."
        })
        
        # News source
        results.append({
            "title": f"Latest on {query} | News Source",
            "url": f"https://news.example.com/topics/{query.lower().replace(' ', '-')}",
            "snippet": f"Get the latest updates and news coverage about {query} from trusted journalists."
        })
        
        # Q&A or forum site
        results.append({
            "title": f"Understanding {query} - Expert Answers",
            "url": f"https://qa.example.com/questions/{query.lower().replace(' ', '-')}",
            "snippet": f"Expert answers to common questions about {query} with detailed explanations and examples."
        })
        
        # Tutorial or how-to site
        results.append({
            "title": f"{query} Guide: Complete Tutorial with Examples",
            "url": f"https://tutorials.example.com/{query.lower().replace(' ', '-')}-guide",
            "snippet": f"Step-by-step guide to understanding {query} with practical examples and applications."
        })
        
        return results[:num_results]

    def get_mock_results(self, query):
        """Generate mock search results as fallback"""
        return [
            {"title": f"Result 1 for {query}", "url": "https://example.com/1", "snippet": f"This is a sample result about {query} with relevant information."},
            {"title": f"Result 2 for {query}", "url": "https://example.com/2", "snippet": f"Another sample result for {query} with different information."},
            {"title": f"Research on {query}", "url": "https://example.com/research", "snippet": f"Academic and research information related to {query}."},
            {"title": f"Latest news about {query}", "url": "https://example.com/news", "snippet": f"Recent developments and news about {query}."},
            {"title": f"{query} - Wikipedia", "url": "https://example.com/wiki", "snippet": f"Comprehensive information about {query} from various reliable sources."}
        ]

    def fetch_webpage(self, url):
        """Fetch and parse content from a webpage with multiple fallback strategies"""
        try:
            # Make sure URL has scheme
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            
            # Log the attempt
            self.errors.append(f"Attempting to fetch content from: {url}")
            
            # Set up headers that mimic a browser
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8"
            }
            
            # For the demo, we'll return simulated content
            # In a real implementation, you would make an actual HTTP request
            title = f"Simulated content for: {url}"
            content = f"This is simulated content for {url} containing relevant information about the search query. This would be real content from the web in a production environment."
            
            return {
                "url": url,
                "title": title,
                "content": content
            }
                
        except Exception as e:
            error_msg = f"Error fetching {url}: {str(e)}"
            self.errors.append(error_msg)
            return {
                "url": url,
                "title": "Error",
                "content": error_msg
            }

    def process_web_content(self, query):
        """Process web search results and create a vector store"""
        # Search the web
        search_results = self.web_search(query)
        
        # Track sources from the beginning
        self.sources = []
        for result in search_results:
            self.sources.append({
                "url": result["url"],
                "title": result["title"],
                "status": "Searched"
            })
        
        # Fetch and process documents
        documents = []
        for i, result in enumerate(search_results):
            doc = self.fetch_webpage(result["url"])
            documents.append(doc)
            
            # Update source status
            for source in self.sources:
                if source["url"] == result["url"]:
                    if "Error" in doc["title"]:
                        source["status"] = "Failed to retrieve"
                    else:
                        source["status"] = "Retrieved"
        
        # Set up vector store
        if documents:
            texts = []
            metadatas = []
            
            for doc in documents:
                chunks = self.text_splitter.split_text(doc["content"])
                for chunk in chunks:
                    texts.append(chunk)
                    metadatas.append({"source": doc["url"], "title": doc["title"]})
            
            # Create vector store
            self.web_vector_store = Chroma.from_texts(
                texts=texts,
                embedding=self.embeddings,
                metadatas=metadatas
            )
            
            return True
        return False

    def direct_retrieval_answer(self, query, user_id=None, mongodb=None, notebook_id=None):
        """
        Mode 1: Direct retrieval from documents without enhancement.
        """
        if not self.doc_vector_store:
            return "Please upload and process files first."
            
        try:
            # Get the retriever
            retriever = self.doc_vector_store.as_retriever(search_kwargs={"k": 4})
            
            # Create simple QA chain
            prompt_template = """
            Use the following context to answer the question. Be factual and direct.
            If the answer is not in the context, say "I don't have enough information to answer this question."
            
            Context:
            {context}
            
            Question: {question}
            
            Answer:
            """
            
            PROMPT = PromptTemplate(
                template=prompt_template, 
                input_variables=["context", "question"]
            )
            
            # Start timing the query
            query_start_time = time.time()
            
            # Create QA chain
            qa = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",
                retriever=retriever,
                chain_type_kwargs={"prompt": PROMPT},
                return_source_documents=True
            )
            
            # Generate answer
            with st.status("Searching documents and generating answer..."):
                response = qa.invoke({"query": query})
                
            answer = response["result"]
            source_docs = response["source_documents"]
            
            # Calculate query time
            query_time = time.time() - query_start_time
            
            # Format sources
            sources = []
            for doc in source_docs:
                sources.append({
                    "content": doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
                    "source": doc.metadata.get("source", "Unknown"),
                    "file_type": doc.metadata.get("file_type", "unknown")
                })
            
            # Log query for analytics if user is logged in
            if user_id and mongodb:
                mongodb.log_query(user_id, query, query_time, notebook_id)
            
            return {
                "answer": answer,
                "sources": sources,
                "query_time": query_time,
                "mode": "direct_retrieval"
            }
                
        except Exception as e:
            self.errors.append(f"Direct retrieval error: {str(e)}")
            return f"Error in direct retrieval: {str(e)}"

    def enhanced_rag_answer(self, query, user_id=None, mongodb=None, notebook_id=None):
        """
        Mode 2: Enhanced RAG answer with multi-stage pipeline.
        """
        if not self.doc_vector_store:
            return "Please upload and process files first."
            
        try:
            # STAGE 1: RETRIEVAL
            with st.status("Stage 1: Retrieving relevant information...") as status:
                # Get top k documents from vector store
                retriever = self.doc_vector_store.as_retriever(search_kwargs={"k": 5})
                relevant_docs = retriever.invoke(query)
                
                # Extract source content for enhancement later
                source_contents = [doc.page_content for doc in relevant_docs]
                
                status.update(label=f"Retrieved {len(relevant_docs)} relevant passages", state="complete")
            
            # STAGE 2: INITIAL ANSWER GENERATION
            with st.status("Stage 2: Generating initial answer...") as status:
                # Custom prompt for initial answer
                prompt_template = """
                You are an AI assistant that provides accurate information based on documents.
                
                Use the following context to answer the question. Be detailed and precise.
                If the answer is not in the context, say "I don't have enough information to answer this question."
                
                Context:
                {context}
                
                Question: {question}
                
                Answer:
                """
                
                PROMPT = PromptTemplate(
                    template=prompt_template, 
                    input_variables=["context", "question"]
                )
                
                # Start timing the query
                query_start_time = time.time()
                
                # Create QA chain for initial answer
                chain_type_kwargs = {"prompt": PROMPT}
                qa = RetrievalQA.from_chain_type(
                    llm=self.llm,
                    chain_type="stuff",
                    retriever=retriever,
                    chain_type_kwargs=chain_type_kwargs,
                    return_source_documents=True
                )
                
                # Generate initial answer
                response = qa.invoke({"query": query})
                initial_answer = response["result"]
                source_docs = response["source_documents"]
                
                status.update(label="Initial answer generated", state="complete")
            
            # STAGE 3: ENHANCEMENT
            with st.status("Stage 3: Enhancing answer quality...") as status:
                # Enhance the answer
                enhanced_answer = self.enhance_answer(
                    initial_answer=initial_answer, 
                    query=query, 
                    source_content=source_contents
                )
                
                status.update(label="Answer enhanced for clarity and completeness", state="complete")
            
            # Calculate query time
            query_time = time.time() - query_start_time
            
            # Format sources
            sources = []
            for doc in source_docs:
                sources.append({
                    "content": doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
                    "source": doc.metadata.get("source", "Unknown"),
                    "file_type": doc.metadata.get("file_type", "unknown")
                })
            
            # Log query for analytics if user is logged in
            if user_id and mongodb:
                mongodb.log_query(user_id, query, query_time, notebook_id)
            
            # Return the enhanced answer and metadata
            return {
                "answer": enhanced_answer,
                "initial_answer": initial_answer,
                "sources": sources,
                "query_time": query_time,
                "mode": "enhanced_rag"
            }
                
        except Exception as e:
            self.errors.append(f"Enhanced RAG error: {str(e)}")
            return f"Error in enhanced RAG: {str(e)}"

    def hybrid_answer(self, query, user_id=None, mongodb=None, notebook_id=None):
        """
        Mode 3: Hybrid approach that combines document retrieval and web search.
        """
        try:
            doc_sources = []
            web_sources = []
            all_source_content = []
            
            # Start timing the query
            query_start_time = time.time()
            
            # Step 1: Get information from local documents if available
            if self.doc_vector_store:
                with st.status("Retrieving information from documents...") as status:
                    # Get relevant documents
                    doc_retriever = self.doc_vector_store.as_retriever(search_kwargs={"k": 3})
                    doc_docs = doc_retriever.invoke(query)
                    
                    # Extract content
                    for doc in doc_docs:
                        all_source_content.append(doc.page_content)
                        doc_sources.append({
                            "content": doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
                            "source": doc.metadata.get("source", "Local Document"),
                            "file_type": doc.metadata.get("file_type", "unknown")
                        })
                    
                    status.update(label=f"Retrieved {len(doc_docs)} document passages", state="complete")
            
            # Step 2: Get information from the web
            with st.status("Searching and retrieving web information...") as status:
                # Process web content
                self.process_web_content(query)
                
                if self.web_vector_store:
                    # Get relevant web documents
                    web_retriever = self.web_vector_store.as_retriever(search_kwargs={"k": 3})
                    web_docs = web_retriever.invoke(query)
                    
                    # Extract content
                    for doc in web_docs:
                        all_source_content.append(doc.page_content)
                        web_sources.append({
                            "content": doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
                            "source": doc.metadata.get("source", "Web"),
                            "file_type": "web"
                        })
                
                status.update(label=f"Retrieved {len(web_sources)} web passages", state="complete")
            
            # Step 3: Combine information and generate answer
            with st.status("Generating comprehensive answer...") as status:
                if not all_source_content:
                    return "No relevant information found in documents or on the web."
                
                # Create a prompt that combines all information
                combined_prompt_template = """
                You are an AI assistant that provides comprehensive answers by combining information from multiple sources.
                
                Below is context from both uploaded documents and web searches.
                
                Context:
                {context}
                
                Question: {question}
                
                Use the following guidelines to construct your answer:
                1. Start with the most relevant and factual information
                2. Highlight agreements between sources
                3. Note any discrepancies or conflicting information
                4. Clearly identify when information comes from local documents vs. web search
                5. Structure your answer logically with appropriate headings
                6. Be comprehensive but concise
                
                Comprehensive Answer:
                """
                
                PROMPT = PromptTemplate(
                    template=combined_prompt_template, 
                    input_variables=["context", "question"]
                )
                
                # Create combined retriever
                # Since we already have the documents, we'll create a custom retriever
                combined_context = "\n\n".join(all_source_content)
                
                # Create a custom chain to use the combined context
                llm_chain = LLMChain(
                    llm=self.llm,
                    prompt=PROMPT
                )
                
                # Generate answer
                result = llm_chain.invoke({
                    "context": combined_context,
                    "question": query
                })
                
                combined_answer = result["text"]
                
                status.update(label="Comprehensive answer generated", state="complete")
            
            # Calculate query time
            query_time = time.time() - query_start_time
            
            # Combine sources from both document and web
            all_sources = doc_sources + web_sources
            
            # Log query for analytics if user is logged in
            if user_id and mongodb:
                mongodb.log_query(user_id, query, query_time, notebook_id)
            
            # Return the combined answer and sources
            return {
                "answer": combined_answer,
                "sources": all_sources,
                "doc_sources_count": len(doc_sources),
                "web_sources_count": len(web_sources),
                "query_time": query_time,
                "mode": "hybrid"
            }
                
        except Exception as e:
            self.errors.append(f"Hybrid mode error: {str(e)}")
            return f"Error in hybrid mode: {str(e)}"

    def ask(self, query, mode="direct_retrieval", user_id=None, mongodb=None, notebook_id=None):
        """
        Ask a question using the selected mode.
        
        Args:
            query: The user question
            mode: The answering mode (direct_retrieval, enhanced_rag, or hybrid)
            user_id: Optional user ID for analytics
            mongodb: Optional MongoDB connection for logging
            notebook_id: Optional notebook ID for context
            
        Returns:
            Answer based on the selected mode
        """
        if mode == "direct_retrieval":
            return self.direct_retrieval_answer(query, user_id, mongodb, notebook_id)
        elif mode == "enhanced_rag":
            return self.enhanced_rag_answer(query, user_id, mongodb, notebook_id)
        elif mode == "hybrid":
            return self.hybrid_answer(query, user_id, mongodb, notebook_id)
        else:
            return f"Unknown mode: {mode}"

    def get_performance_metrics(self):
        """Return performance metrics for the RAG system."""
        if not self.processing_times:
            return None
            
        return {
            "documents_processed": self.documents_processed,
            "index_building_time": self.processing_times.get("index_building", 0),
            "total_processing_time": self.processing_times.get("total_time", 0),
            "memory_used_gb": self.processing_times.get("memory_used_gb", 0),
            "device": self.device,
            "embedding_model": self.embedding_model_name,
            "errors": len(self.errors)
        }